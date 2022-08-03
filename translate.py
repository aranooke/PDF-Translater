from googletrans import Translator;
from docx import Document
from pdfminer.high_level import extract_text
from docx2pdf import convert;
from docx.shared import Pt

document = Document();

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(23)


translator = Translator();
text = extract_text(r"C:\trans\pdftrans\Akunin.pdf",page_numbers=range(3));
arr = text.split(".");
for i in arr:
    try:
        document.add_paragraph(i);
        document.add_paragraph(translator.translate(i,dest = 'ru').text);
    except Exception:
        pass;

print("Programm complete");
document.save("t1.docx");
convert(r"C:\django\t1.docx",r"C:\django\output.pdf");
